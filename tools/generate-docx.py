#!/usr/bin/env python3
"""Generate the printable Word worksheet (ficha de medição e cálculo) for the Engome.

Data (BOM/materials) is dumped from bom.js (the single source of truth shared
with the site) into a JSON file, which this script reads. Reference
illustrations are flattened screenshots of print.html's annotated images
(themselves cropped from Engome Desenho v3.pdf).

Usage: python3 tools/generate-docx.py <bom.json> <iso.png> <topo.png> <output.docx>
"""
import sys
import json
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

bom_json_path, iso_png, topo_png, out_path = sys.argv[1:5]

with open(bom_json_path) as f:
    data = json.load(f)

MATERIALS = data["MATERIALS"]
BOM = data["BOM"]
DIM_INFO = data["DIM_INFO"]

HEADING = RGBColor(0x17, 0x1B, 0x21)
MUTED = RGBColor(0x5B, 0x64, 0x70)

doc = Document()

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.top_margin = Mm(14)
section.bottom_margin = Mm(14)
section.left_margin = Mm(13)
section.right_margin = Mm(13)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_table_borders(table, size=4, color="1B1F24"):
    """Explicit table-wide borders (don't rely on the 'Table Grid' style alone,
    since some renderers only partially honor named table styles)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def set_column_widths(table, widths):
    """Set both the tblGrid (authoritative column layout) and each cell's tcW,
    since some renderers only honor one of the two."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        cols = tblGrid.findall(qn("w:gridCol"))
        for col, width in zip(cols, widths):
            col.set(qn("w:w"), str(width.twips))
    for row in table.rows:
        if len(row.cells) != len(widths):
            continue  # skip rows with merged cells (span already covers full width)
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]


def add_heading(text, size=16, color=HEADING, space_after=2):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(8)
    return p


def add_note(text, center=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = MUTED
    r.italic = True
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# ===================== PAGE 1: ficha de medição =====================
add_heading("Engome — Ficha de medição do bojo")
add_subtitle("Meça o bojo (casco de madeira) conforme ilustrado e preencha os valores abaixo.")

meta = doc.add_table(rows=1, cols=2)
meta.autofit = True
c0, c1 = meta.rows[0].cells
c0.text = "Feito por: ______________________"
c1.text = "Data: ______________________"
doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_picture(iso_png, width=Mm(120))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_note(
    'Ilustração baseada em "Engome Desenho v3.pdf". Dmf (fora) e Dmd (dentro) são medidas na '
    "mesma altura, na cintura onde as arquilhas interna/externa se encontram.",
    center=True,
)

doc.add_paragraph()
tbl = doc.add_table(rows=1 + len(DIM_INFO), cols=4)
tbl.style = "Table Grid"
set_table_borders(tbl)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = False
set_column_widths(tbl, [Mm(18), Mm(65), Mm(32), Mm(35)])
hdr = tbl.rows[0].cells
for i, h in enumerate(["Medida", "O que é", "Valor", "Unid."]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    set_cell_shading(hdr[i], "EEF0F3")
for i, (key, info) in enumerate(DIM_INFO.items(), start=1):
    row = tbl.rows[i].cells
    row[0].text = key
    row[0].paragraphs[0].runs[0].bold = True
    row[1].text = info["label"]
    row[2].text = ""
    row[3].text = "☐ mm    ☐ cm"

doc.add_paragraph()
doc.add_picture(topo_png, width=Mm(55))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_note(
    "Vista de topo — todos os diâmetros são medidos na horizontal, de face a face.",
    center=True,
)

doc.add_paragraph()
notes_hdr = doc.add_paragraph()
r = notes_hdr.add_run("Observações")
r.bold = True
r.font.size = Pt(10)
notes_tbl = doc.add_table(rows=1, cols=1)
notes_tbl.style = "Table Grid"
set_table_borders(notes_tbl)
nc = notes_tbl.rows[0].cells[0]
nc.text = "\n\n\n"

doc.add_page_break()

# ===================== PAGE 2: cálculo + checklist =====================
add_heading("Engome — Cálculo das peças e checklist de fabricação")
add_subtitle(
    "Use as medidas da página anterior nas fórmulas abaixo (ou confira em "
    "calculadora-engome no site) e anote o comprimento de corte de cada peça."
)

calc_tbl = doc.add_table(rows=1, cols=5)
calc_tbl.style = "Table Grid"
set_table_borders(calc_tbl)
calc_tbl.autofit = False
calc_col_widths = [Mm(58), Mm(12), Mm(44), Mm(30), Mm(40)]
hdr = calc_tbl.rows[0].cells
for i, h in enumerate(["Peça", "Qtd.", "Fórmula (mm)", "Resultado unitário", "Total da peça"]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(8.5)
    set_cell_shading(hdr[i], "EEF0F3")

by_material = {}
for item in BOM:
    by_material.setdefault(item["material"], []).append(item)

for mat_key, mat_info in MATERIALS.items():
    items = by_material.get(mat_key, [])
    if not items:
        continue
    row = calc_tbl.add_row().cells
    row[0].merge(row[1]).merge(row[2]).merge(row[3]).merge(row[4])
    row[0].text = mat_info["label"]
    row[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(row[0], "F4F5F7")
    for item in items:
        row = calc_tbl.add_row().cells
        row[0].text = item["label"]
        row[1].text = str(item["qty"])
        row[2].text = item["formula"]
        row[3].text = "____________"
        row[4].text = f"× {item['qty']} = ________"
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)

calc_tbl.autofit = False
set_column_widths(calc_tbl, calc_col_widths)

doc.add_paragraph()
chk_hdr = doc.add_paragraph()
r = chk_hdr.add_run("Checklist de fabricação")
r.bold = True
r.font.size = Pt(13)
add_note("Progresso: ______ / 42 peças concluídas")

check_tbl = doc.add_table(rows=0, cols=3)
check_tbl.style = "Table Grid"
set_table_borders(check_tbl)
row_cells = None
for i, item in enumerate(BOM):
    if i % 3 == 0:
        row_cells = check_tbl.add_row().cells
    cell = row_cells[i % 3]
    p = cell.paragraphs[0]
    r = p.add_run(item["label"] + (f" ({item['qty']}x)" if item["qty"] > 1 else ""))
    r.bold = True
    r.font.size = Pt(8.5)
    items_p = cell.add_paragraph()
    if item["qty"] > 1:
        parts = [f"☐{n}" for n in range(1, item["qty"] + 1)]
        items_p.add_run("  ".join(parts)).font.size = Pt(8.5)
    else:
        items_p.add_run("☐").font.size = Pt(9)

check_tbl.autofit = False
set_column_widths(check_tbl, [Mm(61)] * 3)

doc.add_paragraph()
add_note("Fórmulas e checklist gerados a partir da calculadora online — jtragtenberg.github.io/calculadora-engome")

doc.save(out_path)
print("Wrote", out_path)
